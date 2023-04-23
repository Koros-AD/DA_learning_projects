#8) Задание по библиотекам docx и python-docx (пример итогового файла в demo.docx):
from docx import Document
document=Document()
document.add_heading('Hello,World!',0)
document.add_paragraph('Python Test Paragraph ')
par=document.add_paragraph('This is Python Test Paragraph ')
par.add_run(' which').bold=True
par.add_run(' was created').italic=True
par.add_run(' by writing Python code in Pycharm IDE').underline=True
document.add_heading('Simple Images',level=1)
document.add_heading('Grocery List',level=2)
document.add_paragraph(
    'Cheese',style='List Number'
)
document.add_paragraph(
    'Milk',style='List Number'
)
document.add_paragraph(
    'Sour Cream',style='List Number'
)
document.add_picture('D:\Media/leaf.jpg')
document.save('Test5.docx')